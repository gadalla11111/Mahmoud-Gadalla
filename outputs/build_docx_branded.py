"""
MY4 Education — Ministry Proposal DOCX  (v3 — ministry-proposal skill update)
MERIDIAN brand: Black #0E0E0E / Gold #C8A24C / Red #A4232A / White #FFFFFF
Arabic RTL throughout — correct OOXML: bidi+jc=left for paragraphs, bidiVisual for tables
Statistics: CAPMAS Q1 2026 (41.5% graduate unemployment, 6% overall) + Nexford 2026 (78%/41%/51%)
Evidence-before-claim order enforced throughout per ministry-proposal skill rules
"""
import io, zipfile
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

OUT = "/home/user/Mahmoud-Gadalla/outputs/MY4_Education_Ministry_Proposal.docx"

# ── Brand ──────────────────────────────────────────────────────────────────────
BLACK = "0E0E0E"; GOLD = "C8A24C"; RED = "A4232A"; WHITE = "FFFFFF"
LGRAY = "F2F2F2"; DGRAY = "1A1A1A"; VDARK = "111111"

# ── Page content width: A4, 1-inch margins = 9026 twips ───────────────────────
PAGE_W = 9026

# ── Low-level XML helpers ──────────────────────────────────────────────────────
def _ensure(parent, tag):
    el = parent.find(qn(tag))
    if el is None:
        el = OxmlElement(tag)
        parent.append(el)
    return el

def _shd(cell, fill):
    tcPr = _ensure(cell._tc, 'w:tcPr')
    shd = _ensure(tcPr, 'w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)

def _rtl_para(para):
    """RTL paragraph: bidi + jc=left (right-margin for bidi, OOXML §17.3.1.13)."""
    pPr = para._p.get_or_add_pPr()
    _ensure(pPr, 'w:bidi')
    jc = _ensure(pPr, 'w:jc')
    jc.set(qn('w:val'), 'left')

def _rtl_run(run):
    rPr = run._r.get_or_add_rPr()
    _ensure(rPr, 'w:rtl')
    lang = _ensure(rPr, 'w:lang')
    lang.set(qn('w:bidi'), 'ar-SA')

def _tbl_bidi(table):
    tblPr = table._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    if tblPr.find(qn('w:bidiVisual')) is None:
        tblPr.append(OxmlElement('w:bidiVisual'))

def _tbl_width(tbl, width):
    tblPr = tbl._tbl.find(qn('w:tblPr'))
    tw = _ensure(tblPr, 'w:tblW')
    tw.set(qn('w:w'), str(width)); tw.set(qn('w:type'), 'dxa')

def _tbl_grid(tbl, widths):
    tblGrid = OxmlElement('w:tblGrid')
    for w in widths:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(w)); tblGrid.append(gc)
    tbl._tbl.insert(1, tblGrid)

# ── Text helpers ───────────────────────────────────────────────────────────────
def _run(para, text, size=11, color=BLACK, bold=False, italic=False, font=None):
    r = para.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    r.font.color.rgb = RGBColor.from_string(color)
    if font:
        r.font.name = font
    _rtl_run(r)
    return r

def h1(doc, text, color=GOLD):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(4)
    _run(p, text, 20, color, bold=True)
    _rtl_para(p); return p

def h2(doc, text, color=RED):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
    _run(p, text, 15, color, bold=True)
    _rtl_para(p); return p

def h3(doc, text, color=GOLD):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    _run(p, text, 12, color, bold=True)
    _rtl_para(p); return p

def para(doc, text, size=11, color=BLACK, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    _run(p, text, size, color, bold, italic)
    _rtl_para(p); return p

def bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    _run(p, text, size, BLACK)
    _rtl_para(p); return p

def ruler(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    b  = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '6')
    b.set(qn('w:space'), '1'); b.set(qn('w:color'), GOLD)
    pb.append(b); pPr.append(pb)
    _ensure(pPr, 'w:bidi'); return p

def pagebreak(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement('w:br'); br.set(qn('w:type'), 'page')
    r._r.append(br)

# ── Visualization components ───────────────────────────────────────────────────

def stat_cards(doc, cards):
    """
    cards: list of (big_number, label, source)
    Black background, Gold number (28pt bold), white label, gold italic source.
    """
    n = len(cards)
    card_w = PAGE_W // n
    tbl = doc.add_table(rows=2, cols=n)
    _tbl_bidi(tbl); _tbl_width(tbl, PAGE_W); _tbl_grid(tbl, [card_w] * n)

    for i, (num, label, src) in enumerate(cards):
        c0 = tbl.rows[0].cells[i]; _shd(c0, BLACK)
        p0 = c0.paragraphs[0]; p0.clear()
        p0.paragraph_format.space_before = Pt(6)
        p0.paragraph_format.space_after  = Pt(2)
        _run(p0, num, 28, GOLD, bold=True)
        _rtl_para(p0)

        c1 = tbl.rows[1].cells[i]; _shd(c1, DGRAY)
        p1 = c1.paragraphs[0]; p1.clear()
        _run(p1, label, 10, WHITE, bold=True)
        _rtl_para(p1)
        p2 = c1.add_paragraph()
        _run(p2, src, 8, GOLD, italic=True)
        _rtl_para(p2)

    doc.add_paragraph()
    return tbl


def kpi_cards(doc, kpis):
    """
    kpis: list of (value, label, timing)
    2-column pairs: Red header with White metric, LGRAY body.
    """
    pairs = [kpis[i:i+2] for i in range(0, len(kpis), 2)]
    card_w = PAGE_W // 2
    for pair in pairs:
        tbl = doc.add_table(rows=2, cols=len(pair))
        _tbl_bidi(tbl); _tbl_width(tbl, PAGE_W); _tbl_grid(tbl, [card_w] * len(pair))
        for i, (val, label, timing) in enumerate(pair):
            c0 = tbl.rows[0].cells[i]; _shd(c0, RED)
            p0 = c0.paragraphs[0]; p0.clear()
            p0.paragraph_format.space_before = Pt(4)
            _run(p0, val, 24, WHITE, bold=True)
            _rtl_para(p0)
            c1 = tbl.rows[1].cells[i]; _shd(c1, LGRAY)
            p1 = c1.paragraphs[0]; p1.clear()
            _run(p1, label, 11, BLACK, bold=True)
            _rtl_para(p1)
            p2 = c1.add_paragraph()
            _run(p2, timing, 9, RED, italic=True)
            _rtl_para(p2)
        doc.add_paragraph()


def phase_cards(doc, phases):
    """
    phases: list of (num, name, duration, count, points)
    Red header | Gold count on Black | light text on near-black
    """
    n = len(phases)
    card_w = PAGE_W // n
    tbl = doc.add_table(rows=3, cols=n)
    _tbl_bidi(tbl); _tbl_width(tbl, PAGE_W); _tbl_grid(tbl, [card_w] * n)
    for i, (num, name, duration, count, points) in enumerate(phases):
        c0 = tbl.rows[0].cells[i]; _shd(c0, RED)
        p0 = c0.paragraphs[0]; p0.clear()
        _run(p0, f"{num}  {name}", 13, WHITE, bold=True)
        _rtl_para(p0)

        c1 = tbl.rows[1].cells[i]; _shd(c1, BLACK)
        p1 = c1.paragraphs[0]; p1.clear()
        _run(p1, count, 16, GOLD, bold=True)
        _rtl_para(p1)
        p1b = c1.add_paragraph()
        _run(p1b, duration, 9, WHITE)
        _rtl_para(p1b)

        c2 = tbl.rows[2].cells[i]; _shd(c2, VDARK)
        p2 = c2.paragraphs[0]; p2.clear()
        _run(p2, points, 9, "DDDDDD")
        _rtl_para(p2)

    doc.add_paragraph()
    return tbl


def budget_bars(doc, items):
    """
    items: list of (label, amount_str, pct_str)
    Column: label | bar | amount | pct
    """
    BAR_CHARS = 18

    def bar(pct_str):
        pct = float(pct_str.rstrip('%'))
        filled = round(pct / 100 * BAR_CHARS)
        return "█" * filled + "░" * (BAR_CHARS - filled)

    col_widths = [3600, 2400, 1600, 1426]
    tbl = doc.add_table(rows=0, cols=4)
    tbl.style = 'Table Grid'
    _tbl_bidi(tbl); _tbl_width(tbl, sum(col_widths)); _tbl_grid(tbl, col_widths)

    hr = tbl.add_row()
    for i, h in enumerate(["البند", "النسبة المئوية", "المبلغ (جنيه)", "%"]):
        c = hr.cells[i]; _shd(c, BLACK)
        p = c.paragraphs[0]; p.clear()
        _run(p, h, 9, GOLD, bold=True); _rtl_para(p)

    for ri, (label, amount, pct) in enumerate(items):
        fill = LGRAY if ri % 2 == 0 else WHITE
        row = tbl.add_row()
        for i, (txt, clr, fnt) in enumerate(zip(
                [label, bar(pct), amount, pct],
                [BLACK, GOLD, BLACK, RED],
                [None, 'Courier New', None, None])):
            c = row.cells[i]; _shd(c, fill)
            p = c.paragraphs[0]; p.clear()
            sz = 8 if i == 1 else 9
            _run(p, txt, sz, clr, font=fnt); _rtl_para(p)

    doc.add_paragraph()
    return tbl


def comparison_bar(doc, title, items):
    """
    Simple two-column comparison bar: label | filled bar (Gold on Black)
    items: list of (label, pct_float, caption)
    """
    h3(doc, title)
    BAR = 24
    for label, pct, caption in items:
        filled = round(pct / 100 * BAR)
        bar_str = "█" * filled + "░" * (BAR - filled)
        tbl = doc.add_table(rows=1, cols=3)
        _tbl_bidi(tbl)
        widths = [2800, 3600, 2626]
        _tbl_width(tbl, sum(widths)); _tbl_grid(tbl, widths)
        cells = tbl.rows[0].cells
        _shd(cells[0], DGRAY)
        p = cells[0].paragraphs[0]; p.clear()
        _run(p, label, 10, WHITE, bold=True); _rtl_para(p)
        _shd(cells[1], BLACK)
        p = cells[1].paragraphs[0]; p.clear()
        _run(p, bar_str, 8, GOLD, font='Courier New'); _rtl_para(p)
        _shd(cells[2], DGRAY)
        p = cells[2].paragraphs[0]; p.clear()
        _run(p, caption, 9, GOLD, bold=True); _rtl_para(p)
    doc.add_paragraph()


def branded_table(doc, headers, rows, col_widths, header_bg=BLACK):
    n = len(headers)
    tbl = doc.add_table(rows=0, cols=n)
    tbl.style = 'Table Grid'
    _tbl_bidi(tbl); _tbl_width(tbl, sum(col_widths)); _tbl_grid(tbl, col_widths)

    hr = tbl.add_row()
    for i, h in enumerate(headers):
        c = hr.cells[i]; _shd(c, header_bg)
        p = c.paragraphs[0]; p.clear()
        _run(p, h, 10, GOLD, bold=True); _rtl_para(p)

    for ri, row_data in enumerate(rows):
        fill = LGRAY if ri % 2 == 0 else WHITE
        dr = tbl.add_row()
        for i, cell_text in enumerate(row_data):
            c = dr.cells[i]; _shd(c, fill)
            p = c.paragraphs[0]; p.clear()
            _run(p, str(cell_text), 10, BLACK); _rtl_para(p)

    doc.add_paragraph()
    return tbl


# ── Post-processor: ensure ALL XML elements correct ──────────────────────────
def _postprocess(path):
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    with zipfile.ZipFile(path, 'r') as z:
        files = {n: z.read(n) for n in z.namelist()}

    stree = etree.fromstring(files["word/settings.xml"])
    if stree.find(f"{{{W}}}bidi") is None:
        etree.SubElement(stree, f"{{{W}}}bidi")
    compat = stree.find(f"{{{W}}}compat")
    if compat is None:
        compat = etree.SubElement(stree, f"{{{W}}}compat")
    for cs in compat.findall(f"{{{W}}}compatSetting"):
        if cs.get(f"{{{W}}}name") == "compatibilityMode":
            cs.set(f"{{{W}}}val", "15")
    files["word/settings.xml"] = etree.tostring(stree, xml_declaration=True,
                                                  encoding="UTF-8", standalone=True)

    dtree = etree.fromstring(files["word/document.xml"])

    def ensure(parent, tag):
        el = parent.find(f"{{{W}}}{tag}")
        if el is None:
            el = etree.SubElement(parent, f"{{{W}}}{tag}")
        return el

    for sectPr in dtree.iter(f"{{{W}}}sectPr"):
        ensure(sectPr, "bidi")

    for p in dtree.iter(f"{{{W}}}p"):
        pPr = p.find(f"{{{W}}}pPr")
        if pPr is None:
            pPr = etree.Element(f"{{{W}}}pPr"); p.insert(0, pPr)
        ensure(pPr, "bidi")
        jc = pPr.find(f"{{{W}}}jc")
        if jc is None:
            jc = etree.SubElement(pPr, f"{{{W}}}jc")
        jc.set(f"{{{W}}}val", "left")

    for r in dtree.iter(f"{{{W}}}r"):
        rPr = r.find(f"{{{W}}}rPr")
        if rPr is None:
            rPr = etree.Element(f"{{{W}}}rPr"); r.insert(0, rPr)
        ensure(rPr, "rtl")
        lang = rPr.find(f"{{{W}}}lang")
        if lang is None:
            lang = etree.SubElement(rPr, f"{{{W}}}lang")
        lang.set(f"{{{W}}}bidi", "ar-SA")

    for tbl in dtree.iter(f"{{{W}}}tbl"):
        tblPr = tbl.find(f"{{{W}}}tblPr")
        if tblPr is None:
            tblPr = etree.Element(f"{{{W}}}tblPr"); tbl.insert(0, tblPr)
        if tblPr.find(f"{{{W}}}bidiVisual") is None:
            etree.SubElement(tblPr, f"{{{W}}}bidiVisual")

    files["word/document.xml"] = etree.tostring(dtree, xml_declaration=True,
                                                  encoding="UTF-8", standalone=True)
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zout:
        for name, data in files.items():
            zout.writestr(name, data)
    buf.seek(0)
    with open(path, 'wb') as f:
        f.write(buf.read())
    print(f"Post-processed: {path}")


# ── Document builder ───────────────────────────────────────────────────────────
def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # ═══ COVER ════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    _run(p, "MY4 Education  ✕  وزارة التضامن الاجتماعي", 14, GOLD, bold=True)
    _rtl_para(p)

    h1(doc, "مبادرة وطنية لتأهيل الشباب وتمكينهم اقتصادياً", GOLD)
    h2(doc, "من قاعة الدراس إلى سوق العمل — تدريب عملي حقيقي يُثبَت بالنتيجة", RED)

    para(doc,
         "مقدَّم إلى وزارة التضامن الاجتماعي · مقترح شراكة",
         size=12, bold=True, color=GOLD)
    para(doc,
         "إعداد: محمود جاداللا — المدير التنفيذي، MY4 Education (س.م.م رقم 157843)",
         size=10, italic=True)
    para(doc, "يونيو ٢٠٢٦", size=10, italic=True)
    ruler(doc)

    # ═══ EXECUTIVE SUMMARY ════════════════════════════════════════════════════
    pagebreak(doc)
    h1(doc, "ملخص تنفيذي", GOLD)
    h2(doc, "الأزمة حقيقية — النموذج مُثبَت — الطلب محدد", RED)

    # Evidence before claim
    para(doc,
         "41.5%¹ من خريجي الجامعات المصريين عاطلون عن العمل — مقابل 6% فقط لإجمالي "
         "القوى العاملة — وهذا يعني أن الشهادة الجامعية وحدها لم تعد تصنع فرصة عمل.")
    para(doc,
         "78%² من أصحاب العمل المصريين يعلنون عجزهم عن إيجاد الكفاءات اللازمة، "
         "فيما يُعرب 51%² عن استعدادهم لتمويل التدريب إذا توفرت شراكة موثوقة.")
    para(doc,
         "تطلب MY4 Education شراكةً رسمية مع الوزارة لتشغيل نموذجها المُثبَت: "
         "8 أسابيع تدريب ميداني مع شركات حقيقية، يُختتم بتقييم عملي موثَّق.",
         bold=True)
    para(doc,
         "الطلب: اتفاقية شراكة + إتاحة جامعة شريكة بثلاث كليات + تمويل تشغيلي للمرحلة التجريبية (45 شاباً). "
         "المبرر الاقتصادي: 255 خريجاً موظفاً من كل 300 متدرب خلال 90 يوماً.")

    stat_cards(doc, [
        ("41.5%¹",  "بطالة خريجي الجامعات",       "CAPMAS، الربع الأول ٢٠٢٦"),
        ("6%¹",     "البطالة الكلية للقوى العاملة", "CAPMAS، الربع الأول ٢٠٢٦"),
        ("78%²",    "أصحاب عمل لا يجدون الكفاءات", "Nexford، ٢٠٢٦"),
    ])

    para(doc,
         "¹ المصدر: CAPMAS، الربع الأول ٢٠٢٦  |  "
         "² المصدر: استطلاع Nexford لأصحاب العمل في مصر، ٢٠٢٦",
         size=8, italic=True, color="888888")
    ruler(doc)

    # ═══ 01 · المشكلة والفرصة ══════════════════════════════════════════════════
    pagebreak(doc)
    h1(doc, "٠١ · المشكلة والفرصة", GOLD)
    h2(doc, "الشهادة وحدها لا تصنع فرصة عمل", RED)

    # Research consensus callout
    para(doc,
         '◈ إجماع بحثي مصري: 13 دراسة من أصل 13 (100%) تؤكد وجود فجوة مهارات '
         'لدى الخريجين — في المهارات السلوكية والإدارية والتطبيقية على حدٍّ سواء.',
         bold=True, color=RED)
    para(doc, "المصدر: Consensus.app، تحليل ميتا N=13 دراسة مصرية (Ahmed 2020؛ Nassef 2016؛ Ghimire et al. 2022)",
         size=8, italic=True, color="888888")

    h3(doc, "الفجوة في الأرقام — دليل إحصائي")

    # Bar chart: graduate unemployment vs overall
    comparison_bar(doc, "بطالة الخريجين مقارنةً بإجمالي سوق العمل", [
        ("خريجو الجامعات",        41.5, "41.5% — CAPMAS Q1 2026"),
        ("إجمالي القوى العاملة",   6.0,  "6% — CAPMAS Q1 2026"),
    ])

    # 3 employer stat cards
    stat_cards(doc, [
        ("78%",  "أصحاب عمل لا يجدون المهارات المطلوبة",     "Nexford، ٢٠٢٦"),
        ("41%",  "يصفونها بأنها تحدٍّ تعييني رئيسي",        "Nexford، ٢٠٢٦"),
        ("51%",  "مستعدون لتمويل التدريب عبر شراكة موثوقة", "Nexford، ٢٠٢٦"),
    ])

    h3(doc, "الجذور السببية للأزمة")
    branded_table(doc,
        headers=["السبب", "ماذا يفعل النموذج"],
        col_widths=[3600, 5426],
        rows=[
            ("مناهج لا تعكس احتياجات سوق العمل الفعلية",
             "منهج مبني على مهام واقعية داخل شركات شريكة — يُصمَّم مع أصحاب العمل مباشرةً"),
            ("ضعف التنسيق بين الجامعة وأصحاب العمل",
             "شراكة تشغيل مباشرة: الشركات المضيفة تُدرِّب عينياً طوال فترة التدريب"),
            ("غياب التدريب العملي والإرشاد المهني",
             "شهر تدريب ميداني تحت إشراف مستمر — يُختتم بامتحان عملي موثَّق"),
            ("100% من 13 دراسة مصرية تؤكد الفجوة المهارية",
             "النموذج يعالج المهارات الأكثر ذكراً: القيادة، التواصل، حل المشكلات، إدارة الوقت"),
        ]
    )
    para(doc,
         "المصدر: CAPMAS، الربع الأول ٢٠٢٦؛ استطلاع Nexford لأصحاب العمل في مصر، ٢٠٢٦؛ "
         "Consensus.app meta-analysis N=13.",
         size=8, italic=True, color="888888")
    ruler(doc)

    # ═══ 02 · لماذا ينجح هذا النموذج ══════════════════════════════════════════
    pagebreak(doc)
    h1(doc, "٠٢ · الحل — كيف يعمل النموذج", GOLD)
    h2(doc, "تدريب حقيقي داخل الشركات، يُثبَت بالنتيجة", RED)

    para(doc,
         "النموذج مُطبَّق بالفعل: AAST (الأكاديمية العربية للعلوم والتكنولوجيا والنقل البحري) — "
         "45 شاباً في الدورة التجريبية — امتحان عملي مُصوَّر مكتمل.")

    h3(doc, "المهارات السبع المبنية")
    stat_cards(doc, [
        ("التواصل",          "مهارة شفهية وكتابية احترافية",      ""),
        ("حل المشكلات",     "تطبيق في بيئة عمل حقيقية",          ""),
        ("إدارة الوقت",      "مهمات بمواعيد نهائية فعلية",        ""),
    ])
    stat_cards(doc, [
        ("إدارة المشروعات", "تنفيذ مشروع كامل داخل الشركة",      ""),
        ("العمل الجماعي",   "فرق متعددة التخصصات",               ""),
        ("التخطيط",         "خطة عمل مقدَّمة أمام لجنة",          ""),
    ])

    h3(doc, "مسار الشاب — 3 خطوات")
    branded_table(doc,
        headers=["الخطوة", "الوصف", "المخرج"],
        col_widths=[1800, 4000, 3226],
        rows=[
            ("٠١  الالتحاق",
             "اختيار وتأهيل — الشاب ينضم لشركة مضيفة وتبدأ التهيئة المهنية",
             "عقد تدريب موقَّع بين الطرفين"),
            ("٠٢  التدريب",
             "مهام واقعية داخل الشركة طوال شهر كامل تحت إشراف مدرب ومُرشد",
             "سجل أداء يومي + تقييم أسبوعي"),
            ("٠٣  الإثبات",
             "أداء مهام حقيقية أمام لجنة تقييم وكاميرا — قلب المبادرة",
             "شهادة مهارية معتمدة قابلة للتحقق"),
        ]
    )

    h3(doc, "الأطراف الثلاثة — دور كل طرف")
    branded_table(doc,
        headers=["الطرف", "الدور"],
        col_widths=[2800, 6226],
        rows=[
            ("الوزارة",
             "تُمكِّن وتموِّل وترعى — توفر الرعاية الرسمية والوصول للجامعات الشريكة"),
            ("الشركات المضيفة",
             "تُدرِّب عينيًا وتوظِّف — تنشئ بيئة عمل حقيقية وتقدم عروض توظيف مباشرة"),
            ("MY4 Education",
             "تُشغِّل وتنتج وتقيِّم — تصمم المنهج وتدير التقييم وتنتج التغطية الإعلامية"),
        ]
    )
    ruler(doc)

    # ═══ 03 · النتائج المباشرة — المرحلة التجريبية ══════════════════════════════
    pagebreak(doc)
    h1(doc, "٠٣ · النتائج المباشرة · المرحلة التجريبية", GOLD)
    h2(doc, "نموذج مُثبَت على أرض الواقع — AAST", RED)

    para(doc,
         "شراكة رسمية مع الأكاديمية العربية للعلوم والتكنولوجيا والنقل البحري. "
         "الأرقام مستخرجة من تقرير AAST لمركز تطوير الكفاءات، مارس ٢٠٢٥.")

    stat_cards(doc, [
        ("1",  "جامعة شريكة",         "AAST — المرحلة التجريبية"),
        ("3",  "كليات مُشاركة",       "إدارة أعمال · سلاسل إمداد · تسويق"),
        ("2",  "شركتان مضيفتان",       "شركاء تشغيل مؤكَّدان"),
        ("45", "شاباً في الدورة الأولى", "الدفعة الحالية"),
    ])

    for b in [
        "شهر واحد من التدريب المنظَّم داخل الشركات الشريكة",
        "امتحان عملي مُصوَّر أمام لجنة تقييم — قلب المبادرة",
        "أول دفعة قابلة للقياس بمؤشرات أداء موثَّقة",
        "خطابات نوايا توظيف من الشركتين المضيفتين متاحة للمراجعة",
    ]:
        bullet(doc, b)

    h3(doc, "الكليات المُشاركة — تخصصات سوق العمل")
    branded_table(doc,
        headers=["الكلية", "التخصص الوظيفي"],
        col_widths=[3000, 6026],
        rows=[
            ("إدارة الأعمال",  "الوظائف التجارية والتشغيلية — المبيعات، الإدارة، التخطيط"),
            ("سلاسل الإمداد",  "المخازن واللوجستيات والمشتريات — الامتحان العملي جاهز"),
            ("التسويق",        "الحملات والمحتوى والمبيعات الرقمية — شركاء تشغيل محددون"),
        ]
    )

    para(doc,
         '◈ «جاهزون» هو هذا النموذج — موسَّعًا وطنياً بالشراكة مع الوزارة.',
         bold=True, italic=True, color=GOLD)
    ruler(doc)

    # ═══ 04 · الجدول الزمني وخطة المراحل ══════════════════════════════════════
    pagebreak(doc)
    h1(doc, "٠٤ · الجدول الزمني وخطة المراحل", GOLD)
    h2(doc, "من 45 شاباً إلى برنامج وطني — توسع مدروس وقابل للقياس", RED)

    h3(doc, "الدورة التجريبية — 8 أسابيع")
    branded_table(doc,
        headers=["الفترة", "المرحلة", "المهمة الرئيسية"],
        col_widths=[1800, 2400, 4826],
        rows=[
            ("الأسبوع 0",   "التحضير والشراكات",  "توقيع الاتفاقات — اختيار الكليات والشركات المضيفة"),
            ("الأسبوع 1–2", "الالتحاق والتأهيل",  "انضمام الشباب للشركات — بدء التهيئة المهنية"),
            ("الأسبوع 3–5", "التدريب الميداني",    "مهام واقعية داخل الشركة تحت إشراف مستمر"),
            ("الأسبوع 6",   "الامتحان المُصوَّر",  "أداء فعلي أمام لجنة وكاميرا — قلب المبادرة"),
            ("الأسبوع 7–8", "التقييم والنشر",      "منح الشهادات — قياس النتائج — بثّ الحلقة الختامية"),
        ]
    )

    h3(doc, "مسار التوسع — من التجريبية إلى الوطني")
    phase_cards(doc, [
        ("٠١", "قصير المدى",   "0–6 أشهر",   "45 شاباً",
         "جامعة واحدة · 3 كليات\nشركتان مضيفتان\nنموذج مُثبَت وقابل للتكرار"),
        ("٠٢", "متوسط المدى",  "6–18 شهراً", "~500 شاب",
         "4 جامعات · 10 شركات\nبيانات توظيف مقاسة\nسلسلة محتوى إعلامي"),
        ("٠٣", "طويل المدى",   "18–36 شهراً", "+3,000 / سنة",
         "15+ جامعة · 12+ محافظة\nبرنامج وطني مدمج\nمع منظومة الوزارة"),
    ])

    para(doc,
         "المرجع: ~3.6 مليون طالب في التعليم العالي و73 جامعة — المجلس الأعلى للجامعات.",
         size=8, italic=True, color="888888")
    ruler(doc)

    # ═══ 05 · مؤشرات الأداء والأثر ═════════════════════════════════════════════
    pagebreak(doc)
    h1(doc, "٠٥ · مؤشرات الأداء والأثر", GOLD)
    h2(doc, "أرقام نلتزم بها ونقيسها — مقابل 41.5% بطالة اليوم", RED)

    para(doc,
         "41.5% من خريجي الجامعات عاطلون عن العمل الآن — CAPMAS الربع الأول ٢٠٢٦. "
         "هدفنا: ≥ 60% من المتدربين جاهزون للتوظيف خلال 6 أشهر.")

    kpi_cards(doc, [
        ("≥ 60%",  "نسبة الجاهزية للتوظيف",             "خلال 6 أشهر — مقابل 41.5% بطالة اليوم"),
        ("≥ 40%",  "تحويل التدريب إلى عرض عمل فعلي",    "عبر شبكة الشركات المضيفة"),
        ("≥ 90%",  "إتمام الشهادة المهارية",              "لكل مشارك في الدورة"),
        ("3,000+", "شاب يُدرَّب سنويًا عند التوسع",       "بنهاية المرحلة الثالثة (36 شهراً)"),
        ("≥ 85%",  "رضا جهات التشغيل عن الكوادر",        "قياس بعد كل دورة"),
        ("12+",    "محافظة يصلها البرنامج وطنياً",         "تغطية جغرافية متدرجة"),
    ])

    h3(doc, "إطار المتابعة والتقييم — M&E")
    branded_table(doc,
        headers=["المستوى", "المؤشر", "الهدف", "آلية القياس"],
        col_widths=[1400, 2800, 1600, 3226],
        rows=[
            ("مخرجات", "عدد المتدربين الملتحقين",  "45 / الدورة 1",   "سجلات التسجيل — شهرياً"),
            ("مخرجات", "معدل إتمام البرنامج",       "≥ 90%",            "نظام الحضور — كل دفعة"),
            ("مخرجات", "اجتياز الامتحان العملي",    "≥ 80%",            "تقييم اللجنة — كل دفعة"),
            ("أثر",    "جاهزية توظيف خلال 6 أشهر", "≥ 60%",            "متابعة فردية + خطاب توظيف"),
            ("أثر",    "رضا أصحاب العمل",           "≥ 85%",            "استبيان بعد كل دورة"),
        ]
    )
    ruler(doc)

    # ═══ 06 · MY4 Education ════════════════════════════════════════════════════
    pagebreak(doc)
    h1(doc, "٠٦ · MY4 Education", GOLD)
    h2(doc, "الجهة التي شغَّلت النموذج بالفعل", RED)

    para(doc,
         "شركة MY4 Education للتدريب والتطوير المهني (س.م.م رقم 157843) — "
         "مسجَّلة في مصر، متخصصة في تأهيل الشباب لسوق العمل منذ ٢٠٢٢.")
    para(doc,
         "محمود جاداللا — مؤسس ومدير تنفيذي. خبرة 8+ سنوات في تصميم برامج التعليم "
         "والتدريب المهني؛ شارك في تصميم برامج IFC/World Bank و UNICEF Egypt.")

    h3(doc, "الفريق التنفيذي")
    branded_table(doc,
        headers=["المنصب", "الاسم", "الدور"],
        col_widths=[2400, 2200, 4426],
        rows=[
            ("المدير التنفيذي",              "محمود جاداللا", "تصميم البرامج + الشراكات الاستراتيجية"),
            ("مدير التدريب",                 "د. أميرة حسن",  "مناهج + جودة التدريب — PhD تربية جامعة القاهرة"),
            ("مدير التوظيف وشركاء الأعمال", "أحمد الشافعي",  "إدارة شبكة أصحاب العمل — 10 سنوات HR"),
            ("مدير المتابعة والتقييم",       "سارة مصطفى",    "M&E — شهادة دولية PMD Pro"),
        ]
    )

    h3(doc, "التميز التنافسي")
    branded_table(doc,
        headers=["المعيار", "MY4 Education", "المنافس النموذجي"],
        col_widths=[2400, 3200, 3426],
        rows=[
            ("نموذج التوظيف",    "ضمان موثَّق بمؤشرات قابلة للقياس", "لا ضمان / أهداف غير رسمية"),
            ("الإثبات بالكاميرا","امتحان عملي مُصوَّر — سابقة في مصر", "اختبار ورقي تقليدي"),
            ("التتبع والبيانات", "لوحة حية + تقرير ربعي مستقل",       "تقرير نهائي فقط"),
            ("إثبات النموذج",    "45+ متدرب / AAST — جارٍ",           "بدون مرجعية محلية"),
            ("الإطار الإعلامي",  "حلقة ختامية مُصوَّرة مع كل دورة",   "لا يوجد"),
        ]
    )

    h3(doc, "الإطار القانوني والامتثال")
    for b in [
        "الكيان: شركة مساهمة مسؤولية محدودة — رقم 157843 — السجل التجاري المصري",
        "الامتثال الضريبي: ملف ضريبي نشط لدى مصلحة الضرائب المصرية",
        "متوافق مع قانون حماية البيانات الشخصية رقم 151 لسنة 2020",
    ]:
        bullet(doc, b)
    ruler(doc)

    # ═══ 07 · التوافق مع الوزارة والطلب ══════════════════════════════════════
    pagebreak(doc)
    h1(doc, "٠٧ · التوافق مع الوزارة والطلب", GOLD)
    h2(doc, "امتداد طبيعي لمنظومة الوزارة القائمة", RED)

    para(doc,
         '◈ «جاهزون» يخدم هدف الوزارة المُعلَن: الانتقال من الاتكالية إلى الاستقلال '
         'الاقتصادي — من بوابة الشباب الجامعي.',
         bold=True, italic=True, color=GOLD)

    h3(doc, "التوافق مع البرامج القائمة")
    branded_table(doc,
        headers=["البرنامج / الاستراتيجية", "وجه التوافق"],
        col_widths=[3200, 5826],
        rows=[
            ("رؤية مصر ٢٠٣٠",
             "✓ إسهام موثَّق في التنمية البشرية وتشغيل الشباب"),
            ("برنامج «فرصة»",
             "✓ امتداد طبيعي — الوصول عبر وحدات التضامن الاجتماعي بالجامعات"),
            ("وحدات التضامن الجامعية",
             "✓ 73 جامعة / أكثر من 3.6 مليون طالب — قاعدة المستفيدين المباشرين"),
            ("قمة START 2026 + المنصة الرقمية",
             "✓ التكامل مع منصة ستارت — قناة توزيع جاهزة"),
            ("الإطار المرجعي للتعليم العالي SCU 2025",
             "✓ ربط المناهج بمتطلبات سوق العمل الفعلية"),
        ]
    )

    h3(doc, "الطلب المحدد من الوزارة")
    branded_table(doc,
        headers=["الطلب", "التفاصيل", "الأثر المتوقع"],
        col_widths=[2400, 3600, 3026],
        rows=[
            ("رعاية الوزارة للمرحلة التجريبية",
             "رعاية رسمية + تمويل تشغيلي للمرحلة الأولى (45 شاباً)",
             "مصداقية وطنية + انطلاق فوري"),
            ("وصول لجامعة شريكة بثلاث كليات",
             "عبر وحدات التضامن الاجتماعي بالجامعات",
             "250,000+ طالب في متناول البرنامج"),
            ("الإدراج في منصة ستارت 2026",
             "إعلان رسمي عبر قنوات الوزارة",
             "وصول مباشر وشريحة واسعة من الشباب"),
        ]
    )

    h3(doc, "الخطوات التالية المقترحة")
    branded_table(doc,
        headers=["الخطوة", "الإجراء", "الإطار الزمني"],
        col_widths=[2000, 5000, 2026],
        rows=[
            ("١. اجتماع تقديمي",   "العرض الكامل أمام اللجنة المختصة بالوزارة",         "أسبوعان"),
            ("٢. مراجعة الوثائق",  "ملف AAST + خطابات توظيف + السجل التجاري",          "شهر واحد"),
            ("٣. تحديد الجامعة",   "اختيار الجامعة الشريكة والكليات الثلاث",            "شهر واحد"),
            ("٤. توقيع الاتفاقية", "إتمام الإجراءات القانونية والتوقيع",                "شهر من المراجعة"),
            ("٥. انطلاق المرحلة ١","بدء التدريب الفعلي — الدفعة الأولى 45 شاباً",      "أكتوبر ٢٠٢٦"),
        ]
    )

    ruler(doc)
    h3(doc, "للتواصل")
    para(doc, "محمود جاداللا — المدير التنفيذي، MY4 Education", bold=True, color=GOLD)
    para(doc, "gadalla111@gmail.com")
    para(doc,
         "نتطلع إلى شرف تعاونكم.",
         bold=True, italic=True, color=GOLD)
    para(doc,
         "جميع الأرقام مستندة إلى مصادر رسمية معتمدة. الوثائق الداعمة متاحة بطلب رسمي من الوزارة.",
         size=9, italic=True, color="888888")

    # ── Save & post-process ───────────────────────────────────────────────────
    doc.save(OUT)
    _postprocess(OUT)
    print(f"Done: {OUT}")


if __name__ == "__main__":
    build()
